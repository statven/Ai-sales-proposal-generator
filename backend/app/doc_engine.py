# backend/app/doc_engine.py
import os
import re
import json
import tempfile
import logging
import locale
from io import BytesIO
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches
from docx.table import Table
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont
from datetime import date, datetime # <-- Добавлен импорт date/datetime

DEFAULT_TARGET_DPI = 300
MAX_PAGE_WIDTH_INCHES = 7.3
MAX_PAGE_HEIGHT_INCHES = 8
# --- Импорт безопасных генераторов диаграмм ---
from backend.app.services.visualization_service import (
    generate_gantt_image,
    generate_lifecycle_diagram,
    generate_uml_diagram,
)


logger = logging.getLogger(__name__)

PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")

# --- Валютное форматирование ---
try:
    locale.setlocale(locale.LC_ALL, 'ru_RU.UTF-8')
except locale.Error:
    try:
        locale.setlocale(locale.LC_ALL, 'Russian_Russia')
    except locale.Error:
        logger.warning("Could not set locale for Russian formatting. Using default string conversion.")


def _format_currency(value) -> str:
    if value is None or value == "":
        return ""
    try:
        val = float(value)
        # 2 decimals, thousands separator = space, keep decimal dot (USD)
        s = f"{val:,.2f}"
        # replace comma thousands sep with space (so "10,000.00" -> "10 000.00")
        s = s.replace(",", " ")
        return s
    except Exception:
        return str(value)
def _format_date_english(date_obj: date) -> str:
    """
    Форматирует объект date/datetime в строку на английском языке
    (например, 'December 2, 2025'), игнорируя текущую русскую локаль.
    
    Внимание: Для корректной работы на разных ОС могут понадобиться разные
    идентификаторы локали ('en_US.UTF-8', 'English_United States', 'C').
    """
    if not isinstance(date_obj, (date, datetime)):
        # В случае, если передан не объект даты, возвращаем его строковое представление
        return str(date_obj) 

    # 1. Сохраняем текущую локаль (она должна быть русской)
    original_locale = locale.getlocale(locale.LC_ALL)
    formatted_date = date_obj.isoformat() # Fallback

    # 2. Попытка установить английскую локаль для форматирования
    try:
        # Пытаемся установить локаль 'en_US.UTF-8'
        locale.setlocale(locale.LC_ALL, 'en_US.UTF-8') 
    except locale.Error:
        try:
            # Попытка для Windows/других систем
            locale.setlocale(locale.LC_ALL, 'English_United States')
        except locale.Error:
            # Fallback: Если не удалось, оставляем оригинальную локаль, но используем
            # ISO-формат для английского (YYYY-MM-DD) или другой строгий формат.
            logger.warning("Could not temporarily set English locale for date formatting. Using ISO format.")
            return date_obj.strftime("%B %d, %Y") # Пробуем, но скорее всего будет русский

    # 3. Формат даты на английском языке
    # %B - Full month name (English), %d - Day of month, %Y - Year
    formatted_date = date_obj.strftime("%B %d, %Y")
    
    # 4. Восстанавливаем оригинальную локаль (русскую)
    try:
        locale.setlocale(locale.LC_ALL, original_locale)
    except locale.Error:
        logger.warning("Could not restore original locale after English date formatting.")

    return formatted_date
def sanitize_context(ctx: Dict[str, Any]) -> Dict[str, Any]:
    """
    Clean up context:
    - remove trailing occurrences of client/provider names (signatures, repeated mentions).
    - collapse immediate repeated mentions like "Company X Company X" or "Company X, Company X".
    - do NOT touch special keys in KEEP_KEYS (these are authoritative fields).
    """
    out = dict(ctx)  # shallow copy so we don't modify original directly

    client = (out.get("client_company_name") or out.get("client_name") or "") or ""
    provider = (out.get("provider_company_name") or out.get("provider_name") or "") or ""
    client = str(client).strip()
    provider = str(provider).strip()

    # keys to preserve untouched
    KEEP_KEYS = {
        "client_company_name", "client_name", "client_signature_name",
        "provider_company_name", "provider_name", "provider_signature_name",
        "client_signature_date", "provider_signature_date"
    }

    # helper: remove trailing "ClientName" or "ProviderName" with optional markdown decorators and trailing punctuation/spaces
    def _strip_trailing_names_from_string(s: str, names: List[str]) -> str:

        if not isinstance(s, str) or not s.strip() or not names:
            return s

        # Нормализуем имена для сравнения (приводим к нижнему регистру, убираем лишние пробелы)
        normalized_names = [re.escape(nm.strip()) for nm in names if nm and nm.strip()]

        if not normalized_names:
            return s

        # 1. Удаляем trailing строки, состоящие ТОЛЬКО из имени компании (с любыми markdown/пунктуацией)
        lines = s.split('\n')
        changed = True
        while changed and lines:
            changed = False
            last_line = lines[-1].strip()

            # Удаляем все не-буквенные символы для проверки «это чистое имя?»
            cleaned = re.sub(r'[^\w\s]', '', last_line)  # оставляем только буквы, цифры, пробелы
            cleaned = re.sub(r'\s+', ' ', cleaned).strip()

            if any(re.fullmatch(norm_name, cleaned, flags=re.IGNORECASE) for norm_name in normalized_names):
                lines.pop()
                changed = True
            elif any(cleaned.lower() == re.sub(r'\s+', ' ', nm.lower()) for nm in names if nm):
                lines.pop()
                changed = True

        text = '\n'.join(lines)

        # 2. Удаляем trailing имя в конце всего текста (с markdown, пунктуацией и пробелами)
        pattern = rf'(?:[\s*_\-\.~,.;:!?]*)(?:{"|".join(normalized_names)})(?:[\s*_\-\.~,.;:!?]*)$'
        text = re.sub(pattern, '', text, flags=re.IGNORECASE).strip()

        # 3. Убираем дубликаты имён подряд (например, "Company\nCompany" или "Company, Company")
        for nm in names:
            if not nm:
                continue
            esc = re.escape(nm.strip())
            # Повторения через запятые, точки, пробелы и переносы
            text = re.sub(rf'({esc})\s*[,;:.]?\s*\1', r'\1', text, flags=re.IGNORECASE)

        return text.strip()
    # Apply to each string field except KEEP_KEYS
    for k, v in list(out.items()):
        if k in KEEP_KEYS:
            continue
        if isinstance(v, str):
            out[k] = _strip_trailing_names_from_string(v, [client, provider])

    return out

def _insert_uml_diagram(doc: Document, image_bytes: bytes, placeholder: str, 
                        width_inches: float = 6.5, height_inches: Optional[float] = None):
    """
    Вставляет UML-диаграмму в документ в место плейсхолдера {{uml_diagram}}.
    Если плейсхолдер не найден — добавляет в конец документа.
    """
    inserted = _find_and_replace_placeholder_with_image(doc, placeholder, image_bytes, width_inches, height_inches)
    if not inserted:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        try:
            if height_inches:
                run.add_picture(BytesIO(image_bytes), width=Inches(width_inches), height=Inches(height_inches))
            else:
                run.add_picture(BytesIO(image_bytes), width=Inches(width_inches))
        except Exception:
            p.add_run("[UML image could not be embedded]")

# --- Вспомогательные функции замены текста ---
def _replace_in_paragraph(paragraph, mapping: Dict[str, str]) -> None:
    """
    Replace placeholders like {{key}} inside a paragraph.
    This version:
    - builds full_text from runs
    - performs replacements deterministically (keys by length desc)
    - clears original runs and inserts formatted lines using _apply_formatting_to_run
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return

    # ensure mapping values are strings
    safe_map = {str(k): ("" if v is None else str(v)) for k, v in mapping.items()}

    # quick check for any placeholder existence to avoid unnecessary ops
    if "{{" not in full_text:
        return

    new_text = full_text
    replaced_keys = []
    # replace longer keys first
    for k in sorted(safe_map.keys(), key=lambda x: -len(x)):
        ph = f"{{{{{k}}}}}"
        if ph in new_text:
            new_text = new_text.replace(ph, safe_map[k])
            replaced_keys.append(k)

    if not replaced_keys:
        return

    logger.info("[DOC_ENGINE] Paragraph %s replaced keys: %s", hex(id(paragraph)), replaced_keys)

    # clear existing runs
    for r in paragraph.runs:
        r.text = ""

    # split into lines and re-create runs preserving simple markup
    lines = new_text.split("\n")
    # apply formatting to first line on the same paragraph
    _apply_formatting_to_run(paragraph, lines[0] if lines else "")
    # insert additional paragraphs after current paragraph
    anchor = paragraph._p
    for line in lines[1:]:
        new_p = paragraph._parent.add_paragraph()  # add to same container (section/table cell)
        # preserve paragraph style if any
        try:
            new_p.style = paragraph.style
        except Exception:
            pass
        _apply_formatting_to_run(new_p, line)
        anchor.addnext(new_p._p)
        anchor = new_p._p




def _replace_in_table(table: Table, mapping: Dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _replace_in_paragraph(para, mapping)


def _replace_in_header_footer(container, mapping: Dict[str, str]) -> None:
    if not container:
        return
    for para in container.paragraphs:
        _replace_in_paragraph(para, mapping)
    if hasattr(container, "tables"):
        for t in container.tables:
            _replace_in_table(t, mapping)
def _compute_target_image_inches(png_bytes: bytes,
                                 max_width_in: float = MAX_PAGE_WIDTH_INCHES,
                                 max_height_in: float = MAX_PAGE_HEIGHT_INCHES,
                                 fallback_dpi: int = DEFAULT_TARGET_DPI):
    """
    Возвращает (target_w_in, target_h_in, dpi, orig_w_in, orig_h_in).
    Масштабирует пропорционально, чтобы вписаться в max_width_in x max_height_in.
    Не масштабирует вверх (scale <= 1.0).
    """
    try:
        img = Image.open(BytesIO(png_bytes))
        info = img.info or {}
        dpi = None
        if "dpi" in info:
            dval = info.get("dpi")
            if isinstance(dval, (tuple, list)) and len(dval) >= 1:
                try:
                    dpi = int(dval[0])
                except Exception:
                    dpi = None
            else:
                try:
                    dpi = int(dval)
                except Exception:
                    dpi = None
        if not dpi:
            dpi = int(fallback_dpi)

        orig_w_in = img.width / dpi
        orig_h_in = img.height / dpi

        # если размеры некорректны — fallback
        if orig_w_in <= 0 or orig_h_in <= 0:
            return None, None, dpi, None, None

        # вычисляем масштаб, чтобы вписать в оба ограничения
        scale_w = max_width_in / orig_w_in if orig_w_in > 0 else 1.0
        scale_h = max_height_in / orig_h_in if orig_h_in > 0 else 1.0

        scale = min(scale_w, scale_h, 1.0)  # не увеличиваем (не upscale)

        target_w_in = orig_w_in * scale
        target_h_in = orig_h_in * scale

        return round(target_w_in, 3), round(target_h_in, 3), int(dpi), round(orig_w_in, 3), round(orig_h_in, 3)
    except Exception:
        return None, None, fallback_dpi, None, None


# --- Поиск и вставка изображений ---
def _find_and_replace_placeholder_with_image(doc: Document, placeholder: str, image_bytes: bytes,
                                             width_inches: float = 7.5, height_inches: Optional[float] = None) -> bool:
    """Находит {{placeholder}} и заменяет его изображением"""
    for p in doc.paragraphs:
        if placeholder in p.text:
            for r in list(p.runs):
                r.text = ""
            run = p.add_run()
            try:
                if height_inches:
                    run.add_picture(BytesIO(image_bytes), width=Inches(width_inches), height=Inches(height_inches))
                else:
                    run.add_picture(BytesIO(image_bytes), width=Inches(width_inches))
            except Exception:
                p.add_run("[Image could not be embedded]")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            return True
    # искать в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        for r in list(p.runs):
                            r.text = ""
                        run = p.add_run()
                        try:
                            if height_inches:
                                run.add_picture(BytesIO(image_bytes), width=Inches(width_inches), height=Inches(height_inches))
                            else:
                                run.add_picture(BytesIO(image_bytes), width=Inches(width_inches))
                        except Exception:
                            p.add_run("[Image could not be embedded]")
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        return True
    return False


def _insert_image_with_caption(doc: Document, image_bytes: bytes, placeholder: str,
                               width_inches: float = 8.5, height_inches: Optional[float] = None):
    inserted = _find_and_replace_placeholder_with_image(doc, placeholder, image_bytes, width_inches, height_inches)
    if not inserted:
        # если плейсхолдер не найден — добавляем в конец
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        try:
            if height_inches:
                run.add_picture(BytesIO(image_bytes), width=Inches(width_inches), height=Inches(height_inches))
            else:
                run.add_picture(BytesIO(image_bytes), width=Inches(width_inches))
        except Exception:
            p.add_run("[Image could not be embedded]")
    # подпи



# --- Таблицы ---
def _find_table_by_headers(doc: Document, headers: List[str]) -> Optional[Table]:
    headers_lower = [h.lower() for h in headers]
    for table in doc.tables:
        if not table.rows:
            continue
        first_row = [c.text.strip().lower() for c in table.rows[0].cells]
        if all(any(h in ct for ct in first_row) for h in headers_lower):
            return table
    return None


def _append_deliverables(table: Table, deliverables: List[Dict[str, str]], max_rows: int = 200):
    for d in deliverables[:max_rows]:
        try:
            row = table.add_row()
            cells = row.cells
            title = d.get("title", "")
            desc = d.get("description", "")
            acc = d.get("acceptance_criteria", d.get("acceptance", ""))
            if len(cells) >= 3:
                cells[0].text = title
                cells[1].text = desc
                cells[2].text = acc
            else:
                cells[0].text = f"{title} / {desc} / {acc}"
        except Exception:
            logger.exception("Failed to add deliverable row")


# backend/app/doc_engine.py

def _append_timeline(table: Table, phases: List[Dict[str, Any]], max_rows: int = 200):
    for p in phases[:max_rows]:
        try:
            # Логика приоритетов:
            hours = p.get("duration_hours")
            
            # Если нет hours, ищем duration (иногда приходит как строка "180" или "180 hours")
            if hours is None:
                raw_dur = p.get("duration")
                if raw_dur:
                    try:
                        # Пытаемся вытащить первое число из строки
                        import re
                        nums = re.findall(r'\d+', str(raw_dur))
                        if nums:
                            hours = int(nums[0])
                    except:
                        pass
            
            # Если всё ещё None, ищем недели
            if hours is None and p.get("duration_weeks"):
                try:
                    hours = int(float(p["duration_weeks"]) * 40)
                except:
                    pass

            # Финальный фоллбэк только если совсем ничего нет
            if hours is None:
                hours = 40 

            # Расчет недель для отображения
            weeks = round(hours / 40.0, 1)
            if weeks.is_integer():
                weeks_str = f"{int(weeks)} w"
            else:
                weeks_str = f"{weeks} w"
                
            duration_str = f"{hours} h ({weeks_str})"
            

            name = str(p.get("phase_name", "")).strip()
            tasks = str(p.get("tasks", "")).strip()
            owner = str(p.get("owner", "")).strip()
            priority = str(p.get("priority", "")).strip()

            row = table.add_row()
            cells = row.cells

            if len(cells) >= 3:
                cells[0].text = name
                cells[1].text = duration_str
                if owner:
                    cells[2].text = f"{tasks}\n\nOwner: {owner}"
                else:
                    cells[2].text = tasks
                if priority:
                    cells[2].text += f"\n\nPriority: {priority}"
            else:
                compact = f"{name} — {duration_str}\n{tasks}"
                if owner:
                    compact += f"\nOwner: {owner}"
                if priority:
                    compact += f"\nPriority: {priority}"
                cells[0].text = compact
        except Exception:
            logger.exception("Failed to add timeline row")



def _placeholder_png_bytes(text: str = "Diagram unavailable", width: int = 800, height: int = 400) -> bytes:
    img = Image.new("RGB", (width, height), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 20)
    except IOError:
        font = ImageFont.load_default()
    
    try:
        # draw.textbbox((x, y), text, font) возвращает (left, top, right, bottom)
        bbox = draw.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
    except Exception as e:
        # Fallback values
        logger.error("Error calculating text size using textbbox: %s", str(e))
        text_width, text_height = 300, 25 
    
    # 2. Вычисляем позицию для центрирования
    x = (width - text_width) / 2
    y = (height - text_height) / 2
    
    # 3. Рисуем текст
    draw.text((x, y), text, fill=(50, 50, 50), font=font)
    
    # 4. Сохраняем в буфер
    buf = BytesIO()
    img.save(buf, format="PNG")
    
    return buf.getvalue()

def _insert_lifecycle_diagram(doc: Document, image_bytes: bytes, placeholder: str, 
                              width_inches: float = 6.5, height_inches: Optional[float] = None):
    inserted = _find_and_replace_placeholder_with_image(doc, placeholder, image_bytes, width_inches, height_inches)
    if not inserted:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        try:
            if height_inches:
                run.add_picture(BytesIO(image_bytes), width=Inches(width_inches), height=Inches(height_inches))
            else:
                run.add_picture(BytesIO(image_bytes), width=Inches(width_inches))
        except Exception:
            p.add_run("[Image could not be embedded]")


def render_docx_from_template(template_path: str, context: Dict[str, Any]) -> BytesIO:
    doc = Document(template_path)
    # 1. Prepare mapping (currency formatting preserved)
        # 1. Prepare mapping (currency formatting preserved)
    # Сначала чистим context от лишних подписей
    context = sanitize_context(context)
    today = date.today()
    context['current_date'] = _format_date_english(today)
    mapping = {}
    for k, v in context.items():
        if k in ("development_cost", "licenses_cost", "support_cost", "total_investment_cost"):
            try:
                mapping[k] = _format_currency(v)
            except Exception:
                mapping[k] = str(v or "")
        else:
            mapping[k] = "" if v is None else str(v)

    # diagnostic dump mapping (temporary)
    try:
        dbg = os.path.join(tempfile.gettempdir(), "doc_engine_mapping_dump.json")
        with open(dbg, "w", encoding="utf-8") as fh:
            json.dump(mapping, fh, ensure_ascii=False, indent=2)
        logger.info("[DOC_ENGINE] Mapping dumped to %s", dbg)
    except Exception:
        logger.exception("Failed dumping mapping")



        # 2. Replace placeholders in the document (paragraphs, tables, headers, footers)
    # Мы используем ЕДИНУЮ функцию, которая корректно обрабатывает \n → новые параграфы
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, mapping)

    for table in doc.tables:
        _replace_in_table(table, mapping)

    for section in doc.sections:
        for container in [section.header, section.footer, section.first_page_header, section.first_page_footer]:
            if container:
                for paragraph in container.paragraphs:
                    _replace_in_paragraph(paragraph, mapping)
                for table in container.tables:
                    _replace_in_table(table, mapping)




    # 3. Headers/footers (Эта часть остается без изменений)
    for section in doc.sections:
        try:
            _replace_in_header_footer(section.first_page_header, mapping)
            _replace_in_header_footer(section.first_page_footer, mapping)
            _replace_in_header_footer(section.header, mapping)
            _replace_in_header_footer(section.footer, mapping)
        except Exception:
            logger.exception("Header/footer replacement failed.")
            
    # 4. Tables
    for table in doc.tables:
        try:
            _replace_in_table(table, mapping)
        except Exception:
            logger.exception("Table replace failed.")
            
    # 4.1 Append deliverables & timeline (Эта часть остается без изменений)
    try:
        deliverables_table = _find_table_by_headers(doc, ["Deliverable", "Description", "Acceptance"])
        if deliverables_table and context.get("deliverables_list"):
            _append_deliverables(deliverables_table, context["deliverables_list"])
    except Exception:
        logger.exception("Appending deliverables failed.")

    try:
        timeline_table = _find_table_by_headers(doc, ["Phase", "Duration", "Key Tasks"])
        if timeline_table and context.get("phases_list"):
            _append_timeline(timeline_table, context["phases_list"])
    except Exception:
        logger.exception("Appending timeline failed.")
    # 5. Prepare visualization payload robustly and log it
    def _normalize_visualization_local(ctx: Dict[str, Any]) -> Dict[str, Any]:
        raw = ctx.get("visualization")
        if isinstance(raw, str):
            try:
                parsed = json.loads(raw)
                if isinstance(parsed, dict):
                    raw = parsed
            except Exception:
                raw = {}
        if not isinstance(raw, dict):
            raw = {}

        # fallback keys and permissive mapping for LLM variations
        def pick(d, *keys):
            for k in keys:
                if k in d and d[k]:
                    return d[k]
            return []

        vis = {}
        vis["components"] = pick(raw, "components", "nodes", "elements", "services") or ctx.get("components") or ctx.get("deliverables_list") or []
        vis["data_flows"] = pick(raw, "data_flows", "flows", "edges") or ctx.get("data_flows") or []
        vis["infrastructure"] = pick(raw, "infrastructure", "infra", "servers", "hosts") or ctx.get("infrastructure") or []
        vis["connections"] = pick(raw, "connections", "links", "network") or ctx.get("connections") or []
        vis["milestones"] = pick(raw, "milestones", "timeline", "phases") or ctx.get("phases_list") or ctx.get("milestones") or []
        vis["team_context"] = ctx.get("team_structure_text") or ctx.get("team_structure") or ctx.get("team_roles") or ""
        
        return vis

    viz = _normalize_visualization_local(context)

    # Create diagnostics temp directory (unique per render)
    try:
        tmpdir = tempfile.mkdtemp(prefix="proposal_viz_")
        try:
            with open(os.path.join(tmpdir, "viz_debug.json"), "w", encoding="utf-8") as fh:
                json.dump(viz, fh, ensure_ascii=False, indent=2, default=str)
            logger.info("Visualization payload dumped to %s", os.path.join(tmpdir, "viz_debug.json"))
        except Exception:
            logger.exception("Failed to write viz_debug.json")
    except Exception:
        tmpdir = None
        logger.exception("Failed to create temp dir for viz debugging")

    logger.debug("[DOC_ENGINE] Normalized visualization payload: %s", json.dumps(viz, ensure_ascii=False))

    # 6. Generate lifecycle diagram image
    
    lifecycle_png = None
    try:
        logger.debug("Visualization data for lifecycle diagram: %s", json.dumps(viz, ensure_ascii=False))

        lifecycle_png = generate_lifecycle_diagram(viz)
        if tmpdir and lifecycle_png:
            try:
                with open(os.path.join(tmpdir, "lifecycle.png"), "wb") as fh:
                    fh.write(lifecycle_png)
            except Exception:
                logger.exception("Failed saving lifecycle.png")
        if lifecycle_png and len(lifecycle_png) < 1500:
            logger.warning("Lifecycle image likely placeholder (empty input or export error). Size=%d bytes", len(lifecycle_png))
    except Exception:
        logger.exception("Lifecycle generation raised exception.")

    # 7. Insert images into doc 
    try:

        if tmpdir and lifecycle_png:
            try:
                with open(os.path.join(tmpdir, "lifecycle.png"), "wb") as fh:
                    fh.write(lifecycle_png)
            except Exception:
                logger.exception("Failed saving lifecycle.png")
        if lifecycle_png and len(lifecycle_png) < 1500:
            logger.warning("Lifecycle image likely placeholder (empty input or export error). Size=%d bytes", len(lifecycle_png))

        # Вставка lifecycle: вычисляем реальные инчи и корректируем по странице
                # Вставка lifecycle: вычисляем целевые размеры и корректируем по странице (ширина и высота)
        if lifecycle_png:
            w_in, h_in, dpi, orig_w_in, orig_h_in = _compute_target_image_inches(lifecycle_png)
            if w_in is None:
                # fallback: стандартная вставка
                _insert_lifecycle_diagram(doc, lifecycle_png, "{{lifecycle_diagram}}")
            else:
                logger.debug("Lifecycle image: orig (in) %s x %s @%sdpi -> target (in) %s x %s",
                            orig_w_in, orig_h_in, dpi, w_in, h_in)
                # Передаём ТОЛЬКО width_inches, чтобы сохранить пропорции при вставке.
                # height_inches=None даст python-docx возможность сохранить пропорции.
                _insert_lifecycle_diagram(doc, lifecycle_png, "{{lifecycle_diagram}}",
                                        width_inches=float(w_in), height_inches=None)
        else:
            placeholder_png = _placeholder_png_bytes("Lifecycle diagram unavailable", width=800, height=400)
            _find_and_replace_placeholder_with_image(doc, "{{lifecycle_diagram}}", placeholder_png, width_inches=min(6.5, MAX_PAGE_WIDTH_INCHES), height_inches=None)


    except Exception:
        logger.exception("Inserting lifecycle diagram failed.")

    # 7. Generate diagrams and save intermediate PNGs for debugging
    gantt_png = None
    
    try:
        gantt_png = generate_gantt_image(viz)
        if tmpdir and gantt_png:
            try:
                with open(os.path.join(tmpdir, "gantt.png"), "wb") as fh:
                    fh.write(gantt_png)
            except Exception:
                logger.exception("Failed saving gantt.png")
        if gantt_png and len(gantt_png) < 1500:
            logger.warning("Gantt image likely placeholder (empty input or export error). Size=%d bytes", len(gantt_png))
    except Exception:
        logger.exception("Gantt generation raised exception.")
        # 7.b Generate UML diagram (agent-mode fallback)
    uml_png = None
    try:
        logger.debug("Visualization data for UML diagram: %s", json.dumps(viz, ensure_ascii=False))
        # Call agent-mode UML generator (will fallback deterministically if LLM unavailable)
        uml_png = generate_uml_diagram(viz)
        if tmpdir and uml_png:
            try:
                with open(os.path.join(tmpdir, "uml.png"), "wb") as fh:
                    fh.write(uml_png)
            except Exception:
                logger.exception("Failed saving uml.png")
        if uml_png and len(uml_png) < 1500:
            logger.warning("UML image likely placeholder (empty input or export error). Size=%d bytes", len(uml_png))
    except Exception:
        logger.exception("UML generation raised exception.")

    # Insert UML into doc
    try:
        if uml_png:
            w_in, h_in, dpi, orig_w_in, orig_h_in = _compute_target_image_inches(uml_png)
            if w_in is None:
                # fallback simple insert
                _insert_uml_diagram(doc, uml_png, "{{uml_diagram}}")
            else:
                logger.debug("UML image: orig (in) %s x %s @%sdpi -> target (in) %s x %s",
                            orig_w_in, orig_h_in, dpi, w_in, h_in)
                _insert_uml_diagram(doc, uml_png, "{{uml_diagram}}",
                                    width_inches=float(min(w_in, MAX_PAGE_WIDTH_INCHES)), height_inches=None)
        else:
            placeholder_png = _placeholder_png_bytes("UML diagram unavailable", width=1000, height=420)
            _find_and_replace_placeholder_with_image(doc, "{{uml_diagram}}", placeholder_png, width_inches=min(6.5, MAX_PAGE_WIDTH_INCHES), height_inches=None)
    except Exception:
        logger.exception("Inserting UML diagram failed.")

    # 8. Insert images into doc (if present). If missing, insert explicit note


    try:
        if gantt_png:
            w_in, h_in, dpi, orig_w_in, orig_h_in = _compute_target_image_inches(gantt_png)
            if w_in is None:
                _insert_image_with_caption(doc, gantt_png, "{{gantt_chart}}")
            else:
                logger.debug("Gantt image: orig (in) %s x %s @%sdpi -> target (in) %s x %s",
                            orig_w_in, orig_h_in, dpi, w_in, h_in)
                _insert_image_with_caption(doc, gantt_png, "{{gantt_chart}}",
                                        width_inches=float(w_in), height_inches=None)
        else:
            placeholder_png = _placeholder_png_bytes("Gantt chart unavailable", width=1000, height=500)
            _find_and_replace_placeholder_with_image(doc, "{{gantt_chart}}", placeholder_png, width_inches=min(7.5, MAX_PAGE_WIDTH_INCHES), height_inches=None)


    except Exception:
        logger.exception("Inserting gantt failed.")

    # 9. Save docx to BytesIO and also log location of debug dir (if any)
    out = BytesIO()
    try:
        doc.save(out)
        out.seek(0)
    except Exception:
        logger.exception("Failed saving docx to BytesIO.")
        raise

    if tmpdir:
        logger.info("Visualization debug files saved to: %s", tmpdir)
    return out

def _apply_formatting_to_run(paragraph, text_line: str):
    """
    Apply basic markdown-like formatting for **bold** and *italic*.
    Handles multiple occurrences in the same line.
    Does not attempt to parse nested or malformed markup.
    """
    if text_line is None:
        text_line = ""
    text_line = str(text_line)

    # pattern finds either **bold** or *italic*; non-greedy.
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*)")

    last_idx = 0
    for m in pattern.finditer(text_line):
        start, end = m.span()
        # plain text before match
        if start > last_idx:
            paragraph.add_run(text_line[last_idx:start])
        # determine which group matched
        bold_text = m.group(2)
        italic_text = m.group(3)
        if bold_text is not None:
            r = paragraph.add_run(bold_text)
            r.bold = True
        elif italic_text is not None:
            r = paragraph.add_run(italic_text)
            r.italic = True
        last_idx = end

    # tail
    if last_idx < len(text_line):
        paragraph.add_run(text_line[last_idx:])
